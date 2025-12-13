"""
任务API路由
"""

from typing import Optional, Dict, List
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
import logging
import zipfile
import io
import os
import asyncio
import json
from datetime import datetime, timedelta, timezone

from ..database import get_db
from ..middleware.auth import get_current_user
from ..models.user import User
from ..models.task import Task, TaskPolicy
from ..models.policy import Policy
from ..schemas.task import TaskCreate, TaskResponse, TaskListItem, TaskListResponse
from ..services.task_service import TaskService
from ..services.storage_service import StorageService

router = APIRouter(prefix="/tasks", tags=["tasks"])
task_service = TaskService()
logger = logging.getLogger(__name__)


def _generate_markdown_from_policy(policy) -> str:
    """从政策对象生成Markdown内容"""
    md_lines = []

    # YAML Front Matter
    md_lines.append("---")
    md_lines.append(f'title: "{policy.title}"')
    if policy.level:
        md_lines.append(f'level: "{policy.level}"')
    if policy.category:
        md_lines.append(f'category: "{policy.category}"')
    if policy.pub_date:
        md_lines.append(f'pub_date: "{policy.pub_date}"')
    if policy.doc_number:
        md_lines.append(f'doc_number: "{policy.doc_number}"')
    if policy.effective_date:
        md_lines.append(f'effective_date: "{policy.effective_date}"')
    if policy.validity:
        md_lines.append(f'validity: "{policy.validity}"')
    if policy.publisher:
        md_lines.append(f'publisher: "{policy.publisher}"')
    if policy.source_url:
        md_lines.append(f'source_url: "{policy.source_url}"')
    md_lines.append("---")
    md_lines.append("")

    md_lines.append(f"# {policy.title}")
    md_lines.append("")

    md_lines.append("## 基本信息")
    md_lines.append("")
    if policy.publisher:
        md_lines.append(f"- **发布机构**: {policy.publisher}")
    if policy.level:
        md_lines.append(f"- **效力级别**: {policy.level}")
    if policy.pub_date:
        md_lines.append(f"- **发布日期**: {policy.pub_date}")
    if policy.doc_number:
        md_lines.append(f"- **发文字号**: {policy.doc_number}")
    if policy.effective_date:
        md_lines.append(f"- **生效日期**: {policy.effective_date}")
    if policy.validity:
        md_lines.append(f"- **有效性**: {policy.validity}")
    if policy.category:
        md_lines.append(f"- **分类**: {policy.category}")
    if policy.source_url:
        md_lines.append(f"- **来源链接**: [查看原文]({policy.source_url})")
    md_lines.append("")

    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 正文内容")
    md_lines.append("")
    if policy.content:
        md_lines.append(policy.content)
    else:
        md_lines.append("> **注意**: 该政策的正文内容无法自动获取。")
        md_lines.append("> ")
        md_lines.append("> 请访问[来源链接](#基本信息)查看完整文档内容。")

    return "\n".join(md_lines)


def _generate_docx_from_policy(policy, output_path: str, converter):
    """从政策对象生成DOCX文件"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx未安装，无法生成DOCX")
        return

    # 创建文档
    doc = Document()

    # 设置文档标题
    title = doc.add_heading(policy.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加基本信息
    doc.add_heading("基本信息", level=2)
    info_para = doc.add_paragraph()

    if policy.publisher:
        info_para.add_run("发布机构: ").bold = True
        info_para.add_run(policy.publisher)
        info_para.add_run("\n")
    if policy.level:
        info_para.add_run("效力级别: ").bold = True
        info_para.add_run(policy.level)
        info_para.add_run("\n")
    if policy.pub_date:
        info_para.add_run("发布日期: ").bold = True
        info_para.add_run(str(policy.pub_date))
        info_para.add_run("\n")
    if policy.doc_number:
        info_para.add_run("发文字号: ").bold = True
        info_para.add_run(policy.doc_number)
        info_para.add_run("\n")
    if policy.effective_date:
        info_para.add_run("生效日期: ").bold = True
        info_para.add_run(str(policy.effective_date))
        info_para.add_run("\n")
    if policy.validity:
        info_para.add_run("有效性: ").bold = True
        info_para.add_run(policy.validity)
        info_para.add_run("\n")
    if policy.category:
        info_para.add_run("分类: ").bold = True
        info_para.add_run(policy.category)
        info_para.add_run("\n")

    # 添加正文内容
    doc.add_heading("正文内容", level=2)
    if policy.content:
        # 将内容按段落分割
        paragraphs = policy.content.split("\n")
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    else:
        doc.add_paragraph(
            "该政策的正文内容无法自动获取，请访问来源链接查看完整文档内容。"
        )

    # 保存文档
    doc.save(output_path)


@router.post("/", response_model=TaskResponse, status_code=201)
def create_task(
    task_data: TaskCreate,
    auto_start: bool = Query(True, description="是否自动启动任务"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """创建任务"""
    try:
        task = task_service.create_task(
            db=db,
            task_name=task_data.task_name,
            task_type=task_data.task_type,
            config=task_data.config,
            user_id=current_user.id,
        )

        # 如果设置了自动启动，立即启动任务
        if auto_start:
            try:
                task = task_service.start_task(db, task.id, background=True)
            except Exception as e:
                logger.warning(f"自动启动任务失败: {e}")

        # 在序列化前确保所有字段都已加载（避免延迟加载问题）
        _ = task.id, task.task_name, task.task_type, task.status, task.created_at

        return TaskResponse.model_validate(task)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"创建任务失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")


@router.get("/", response_model=TaskListResponse)
def get_tasks(
    skip: int = Query(0, ge=0, description="跳过数量"),
    limit: int = Query(20, ge=1, le=100, description="每页数量"),
    task_type: Optional[str] = Query(None, description="任务类型筛选"),
    status: Optional[str] = Query(None, description="状态筛选"),
    completed_only: bool = Query(False, description="只返回已完成的任务"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取任务列表"""
    try:
        # 如果请求只返回已完成的任务，设置status参数
        actual_status = status
        if completed_only:
            actual_status = "completed"

        tasks, total = task_service.get_tasks(
            db=db,
            skip=skip,
            limit=limit,
            task_type=task_type,
            status=actual_status,
            completed_only=completed_only,
        )

        # 在序列化前确保所有字段都已加载
        items = []
        for task in tasks:
            try:
                # 访问所有需要的字段以触发加载
                _ = (
                    task.id,
                    task.task_name,
                    task.task_type,
                    task.status,
                    task.created_at,
                )
                item = TaskListItem.model_validate(task)
                items.append(item)
            except Exception as e:
                logger.warning(f"序列化任务失败 (ID: {task.id}): {e}")
                continue

        return TaskListResponse(items=items, total=total, skip=skip, limit=limit)
    except Exception as e:
        logger.error(f"获取任务列表失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取任务列表失败: {str(e)}")


@router.get("/{task_id}", response_model=TaskResponse)
def get_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取任务详情"""
    task = task_service.get_task(db, task_id)

    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    return TaskResponse.model_validate(task)


@router.post("/{task_id}/start", response_model=TaskResponse)
def start_task(
    task_id: int,
    background: bool = Query(True, description="是否后台执行"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """启动任务"""
    try:
        task = task_service.start_task(db, task_id, background=background)
        return TaskResponse.model_validate(task)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"启动任务失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"启动任务失败: {str(e)}")


@router.post("/{task_id}/stop", response_model=TaskResponse)
def stop_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """停止/取消任务"""
    try:
        success = task_service.stop_task(db, task_id)
        if not success:
            raise HTTPException(
                status_code=400, detail="任务无法停止（可能不在运行状态）"
            )

        # 获取更新后的任务对象
        task = db.query(Task).filter(Task.id == task_id).first()
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")

        return TaskResponse.model_validate(task)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"停止任务失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"停止任务失败: {str(e)}")


@router.post("/{task_id}/pause", response_model=TaskResponse)
def pause_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """暂停任务"""
    try:
        success = task_service.pause_task(db, task_id)
        if not success:
            raise HTTPException(
                status_code=400, detail="任务无法暂停（可能不在运行状态）"
            )

        # 获取更新后的任务对象
        task = db.query(Task).filter(Task.id == task_id).first()
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")

        return TaskResponse.model_validate(task)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"暂停任务失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"暂停任务失败: {str(e)}")


@router.post("/{task_id}/resume", response_model=TaskResponse)
def resume_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """恢复任务"""
    try:
        success = task_service.resume_task(db, task_id)
        if not success:
            raise HTTPException(
                status_code=400, detail="任务无法恢复（可能不在暂停状态）"
            )

        # 获取更新后的任务对象
        task = db.query(Task).filter(Task.id == task_id).first()
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")

        return TaskResponse.model_validate(task)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"恢复任务失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"恢复任务失败: {str(e)}")


@router.delete("/{task_id}")
def delete_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """删除任务"""
    try:
        task_service.delete_task(db, task_id)
        return {"message": "任务已删除", "id": task_id}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"删除任务失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"删除任务失败: {str(e)}")


@router.get("/{task_id}/download")
def download_task_files(
    task_id: int,
    file_format: str = Query("all", description="文件格式: all, markdown, docx"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载任务的所有文件（打包成zip）

    Args:
        task_id: 任务ID
        file_format: 文件格式筛选 (all, markdown, docx)

    Returns:
        ZIP文件流
    """
    try:
        # 获取任务
        task = db.query(Task).filter(Task.id == task_id).first()
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")

        # 获取任务关联的所有政策
        task_policies = db.query(TaskPolicy).filter(TaskPolicy.task_id == task_id).all()
        if not task_policies:
            raise HTTPException(status_code=404, detail="该任务没有关联的政策")

        policy_ids = [tp.policy_id for tp in task_policies]
        policies = db.query(Policy).filter(Policy.id.in_(policy_ids)).all()

        if not policies:
            raise HTTPException(status_code=404, detail="未找到任何政策")

        # 初始化存储服务
        storage_service = StorageService()

        # 对于大量文件，使用临时文件而不是内存缓冲区
        # 创建临时ZIP文件路径
        import tempfile

        temp_zip_fd, temp_zip_path = tempfile.mkstemp(suffix=".zip")
        os.close(temp_zip_fd)  # 关闭文件描述符，让zipfile使用路径

        try:
            # 使用 ZIP_STORED 存储模式（不压缩）以提升速度
            # 对于文本文件，压缩效果有限，但速度慢
            # 如果需要压缩，可以改为 ZIP_DEFLATED，但压缩级别设为最低
            with zipfile.ZipFile(temp_zip_path, "w", zipfile.ZIP_STORED) as zip_file:
                file_count = 0

                # 确定要下载的文件类型
                file_types = []
                if file_format == "all":
                    file_types = ["markdown", "docx"]
                elif file_format == "markdown":
                    file_types = ["markdown"]
                elif file_format == "docx":
                    file_types = ["docx"]
                else:
                    raise HTTPException(
                        status_code=400, detail=f"不支持的文件格式: {file_format}"
                    )

                # 导入文件生成工具
                from ..core.converter import DocumentConverter

                converter = DocumentConverter()

                # 遍历每个政策，获取文件 - 添加进度日志
                total_policies = len(policies)
                logger.info(f"开始打包 {total_policies} 个政策的文件")

                for i, policy in enumerate(policies, 1):
                    policy_title = (
                        policy.title.replace("/", "_")
                        .replace("\\", "_")
                        .replace(":", "_")
                    )
                    safe_title = "".join(
                        c
                        for c in policy_title
                        if c.isalnum() or c in (" ", "-", "_", "(", ")", "（", "）")
                    )[:100]

                    for file_type in file_types:
                        # 直接构建本地存储路径（不通过get_policy_file_path，因为它优先返回缓存路径）
                        file_ext = "md" if file_type == "markdown" else file_type
                    if policy.task_id:
                        file_dir = f"{policy.task_id}/{policy.id}"
                    else:
                        file_dir = str(policy.id)

                    file_path = os.path.join(
                        str(storage_service.local_dir),
                        "policies",
                        file_dir,
                        f"{policy.id}.{file_ext}",
                    )

                    # 检查本地存储文件是否存在
                    if os.path.exists(file_path):
                        # 文件存在，直接使用
                        logger.debug(f"使用现有文件: {file_path}")
                    else:
                        # 文件不存在，尝试重新生成
                        logger.warning(
                            f"政策 {policy.id} 的 {file_type} 文件不存在，尝试重新生成..."
                        )
                        try:
                            # 使用storage_service的临时目录，确保文件可以被多个任务共享
                            from ..services.file_cleanup_service import (
                                get_cleanup_service,
                            )

                            cleanup_service = get_cleanup_service()

                            # 创建临时目录（基于policy_id，确保同一政策的文件可以被多个任务共享）
                            temp_base_dir = (
                                storage_service.local_dir
                                / "temp_generated"
                                / str(policy.id)
                            )
                            temp_base_dir.mkdir(parents=True, exist_ok=True)
                            # 将文件类型转换为实际扩展名
                            temp_file_ext = (
                                "md" if file_type == "markdown" else file_type
                            )
                            temp_file_path = (
                                temp_base_dir / f"{policy.id}.{temp_file_ext}"
                            )

                            # 如果临时文件已存在且未过期（1天内），直接使用
                            if temp_file_path.exists():
                                file_stat = temp_file_path.stat()
                                file_age = datetime.now(
                                    timezone.utc
                                ) - datetime.fromtimestamp(
                                    file_stat.st_mtime, tz=timezone.utc
                                )
                                if file_age < timedelta(hours=24):
                                    file_path = str(temp_file_path)
                                    logger.debug(f"使用已存在的临时文件: {file_path}")
                                else:
                                    # 删除过期文件
                                    temp_file_path.unlink()
                            else:
                                # 根据文件类型生成文件
                                if file_type == "markdown":
                                    # 生成Markdown文件
                                    md_content = _generate_markdown_from_policy(policy)
                                    with open(
                                        temp_file_path, "w", encoding="utf-8"
                                    ) as f:
                                        f.write(md_content)
                                elif file_type == "docx":
                                    # 生成DOCX文件
                                    _generate_docx_from_policy(
                                        policy, str(temp_file_path), converter
                                    )

                                # 如果生成成功，注册临时文件并保存到storage_service
                                if (
                                    temp_file_path.exists()
                                    and temp_file_path.stat().st_size > 0
                                ):
                                    file_path = str(temp_file_path)
                                    # 注册临时文件，用于后续清理
                                    cleanup_service.register_temp_file(file_path)

                                    # 尝试保存到storage_service（使用policy的task_id，确保文件路径正确）
                                    try:
                                        storage_result = (
                                            storage_service.save_policy_file(
                                                policy.id,
                                                file_type,
                                                file_path,
                                                task_id=policy.task_id,
                                            )
                                        )
                                        if storage_result.get("success"):
                                            # 如果保存成功，使用正式存储路径
                                            file_path = storage_result.get("local_path")
                                            logger.debug(
                                                f"文件已保存到存储服务: {file_path}"
                                            )
                                    except Exception as e:
                                        logger.warning(f"保存文件到存储服务失败: {e}")

                                    logger.info(
                                        f"成功重新生成政策 {policy.id} 的 {file_type} 文件"
                                    )
                                else:
                                    logger.warning(
                                        f"重新生成政策 {policy.id} 的 {file_type} 文件失败"
                                    )
                                    continue
                        except Exception as e:
                            logger.error(
                                f"重新生成政策 {policy.id} 的 {file_type} 文件时出错: {e}"
                            )
                            continue

                    if file_path and os.path.exists(file_path):
                        # 构造ZIP内的文件路径
                        # 格式: 任务名称/文件格式/政策ID_标题.扩展名
                        # 将文件类型转换为实际扩展名
                        file_ext = "md" if file_type == "markdown" else file_type
                        zip_path = f"{task.task_name}/{file_type}/{policy.id}_{safe_title}.{file_ext}"

                        # 添加到ZIP文件 - 优化文件读取
                        try:
                            # 使用二进制模式读取文件
                            with open(file_path, "rb") as f:
                                file_data = f.read()
                                zip_file.writestr(zip_path, file_data)
                            file_count += 1
                            logger.debug(
                                f"添加文件到ZIP: {zip_path} ({len(file_data)} bytes)"
                            )
                        except Exception as e:
                            logger.warning(f"读取文件失败 {file_path}: {e}")
                            continue

                        # 每处理10个政策记录一次进度
                        if i % 10 == 0 or i == total_policies:
                            logger.info(
                                f"已处理 {i}/{total_policies} 个政策，当前文件数: {file_count}"
                            )

                if file_count == 0:
                    raise HTTPException(
                        status_code=404, detail="未找到任何可下载的文件"
                    )

        except Exception as zip_error:
            # 确保临时文件被清理
            try:
                if "temp_zip_path" in locals() and os.path.exists(temp_zip_path):
                    os.unlink(temp_zip_path)
                    logger.debug(f"已清理临时ZIP文件: {temp_zip_path}")
            except Exception as cleanup_error:
                logger.warning(f"清理临时ZIP文件失败: {cleanup_error}")

            # 重新抛出原始异常
            raise zip_error

        # 获取文件大小
        final_size = os.path.getsize(temp_zip_path)
        logger.info(f"ZIP文件打包完成: {file_count} 个文件, 大小: {final_size} bytes")

        # 检查文件大小，如果为0则报错
        if final_size == 0:
            logger.error(
                f"ZIP文件大小为0，可能是打包过程中出现问题。临时文件路径: {temp_zip_path}"
            )
            raise HTTPException(
                status_code=500,
                detail="文件打包失败，生成的文件大小为0字节。请检查系统内存和磁盘空间。",
            )

        # 生成文件名 - 使用ASCII安全的文件名避免编码问题
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        # 将任务名称转换为ASCII安全的格式
        safe_task_name = "".join(c if ord(c) < 128 else "_" for c in task.task_name)
        zip_filename = f"{safe_task_name}_{timestamp}.zip"

        # 创建文件流生成器，用于StreamingResponse
        def file_generator():
            try:
                with open(temp_zip_path, "rb") as f:
                    while chunk := f.read(8192):  # 每次读取8KB
                        yield chunk
            finally:
                # 在生成器完成后删除临时文件
                try:
                    os.unlink(temp_zip_path)
                    logger.debug(f"已删除临时ZIP文件: {temp_zip_path}")
                except Exception as e:
                    logger.warning(f"删除临时ZIP文件失败: {e}")

        # 返回ZIP文件流 - 使用更兼容的Content-Disposition格式
        return StreamingResponse(
            file_generator(),
            media_type="application/zip",
            headers={
                "Content-Disposition": f'attachment; filename="{zip_filename}"',
                "Content-Length": str(final_size),
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载任务文件失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"下载任务文件失败: {str(e)}")


# SSE 连接管理器
class TaskProgressManager:
    """任务进度SSE管理器"""

    def __init__(self):
        self.active_connections: Dict[int, List[asyncio.Queue]] = (
            {}
        )  # task_id -> [queues]

    async def connect(self, task_id: int) -> asyncio.Queue:
        """建立SSE连接"""
        if task_id not in self.active_connections:
            self.active_connections[task_id] = []

        queue = asyncio.Queue()
        self.active_connections[task_id].append(queue)
        return queue

    def disconnect(self, task_id: int, queue: asyncio.Queue):
        """断开SSE连接"""
        if task_id in self.active_connections:
            try:
                self.active_connections[task_id].remove(queue)
            except ValueError:
                pass
            # 如果没有连接了，清理
            if not self.active_connections[task_id]:
                del self.active_connections[task_id]

    async def broadcast_progress(self, task_id: int, data: dict):
        """广播进度更新"""
        if task_id not in self.active_connections:
            return

        # 准备SSE数据
        sse_data = f"data: {json.dumps(data)}\n\n"

        # 发送给所有连接的客户端
        disconnected_queues = []
        for queue in self.active_connections[task_id]:
            try:
                await queue.put(sse_data)
            except Exception:
                # 连接可能已断开
                disconnected_queues.append(queue)

        # 清理断开的连接
        for queue in disconnected_queues:
            self.disconnect(task_id, queue)


# 全局进度管理器实例
progress_manager = TaskProgressManager()


async def generate_progress_events(task_id: int, db: Session):
    """生成SSE事件流"""
    logger.info(f"建立SSE连接: task_id={task_id}")
    try:
        # 建立连接
        queue = await progress_manager.connect(task_id)
        logger.info(f"SSE队列已创建: task_id={task_id}")

        # 发送初始任务状态
        task = task_service.get_task(db, task_id)
        if task:
            initial_data = {
                "type": "task_update",
                "task_id": task.id,
                "status": task.status,
                "progress_message": task.progress_message or "",
                "start_time": task.start_time.isoformat() if task.start_time else None,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }
            logger.info(f"发送初始任务状态: {task.status}")
            yield f"data: {json.dumps(initial_data)}\n\n"

        # 发送连接确认消息
        connection_data = {
            "type": "connection_established",
            "task_id": task_id,
            "message": "SSE连接已建立",
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }
        logger.info(f"发送连接确认消息")
        yield f"data: {json.dumps(connection_data)}\n\n"

        # 发送测试进度消息
        test_data = {
            "type": "progress_update",
            "task_id": task_id,
            "message": "测试进度更新消息",
            "progress_message": "SSE连接测试消息\n连接已建立，等待任务开始...",
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
        logger.info(f"发送测试进度消息")
        yield f"data: {json.dumps(test_data)}\n\n"

        # 持续监听进度更新
        message_count = 0
        while True:
            try:
                # 等待进度更新或超时
                sse_data = await asyncio.wait_for(queue.get(), timeout=30.0)
                message_count += 1
                logger.info(f"SSE消息 {message_count}: 发送进度更新")
                yield sse_data
            except asyncio.TimeoutError:
                # 30秒心跳
                heartbeat_data = {
                    "type": "heartbeat",
                    "message": "连接正常",
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                }
                logger.info(f"SSE心跳: task_id={task_id}")
                yield f"data: {json.dumps(heartbeat_data)}\n\n"
            except Exception as e:
                logger.error(f"SSE监听异常: {e}")
                break

    except Exception as e:
        logger.error(f"SSE连接错误: {e}")
        # 发送错误消息
        error_data = {
            "type": "error",
            "message": f"连接错误: {str(e)}",
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }
        yield f"data: {json.dumps(error_data)}\n\n"
    finally:
        # 清理连接
        logger.info(f"清理SSE连接: task_id={task_id}")
        progress_manager.disconnect(task_id, queue)


@router.get("/{task_id}/progress/stream")
async def stream_task_progress(
    task_id: int,
    token: Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    """流式推送任务进度更新（SSE）"""
    # 验证token
    if not token:
        raise HTTPException(status_code=401, detail="缺少认证令牌")

    from ..services.auth_service import AuthService

    current_user = AuthService.get_current_user(db, token)
    if not current_user:
        raise HTTPException(status_code=401, detail="无效的认证令牌")

    # 检查任务是否存在且属于当前用户
    task = task_service.get_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    if task.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="无权访问此任务")

    return StreamingResponse(
        generate_progress_events(task_id, db),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Cache-Control",
        },
    )
