# ==============================================================================
# MNR Law Crawler Online - 爬虫核心模块
# ==============================================================================
#
# 项目名称: MNR Law Crawler Online (自然资源部法规爬虫系统 - Web版)
# 项目地址: https://github.com/ViVi141/MNR-Law-Crawler-Online
# 作者: ViVi141
# 许可证: MIT License
#
# 描述: 爬虫核心模块，负责政策爬取的主逻辑
#       适配自然资源部政府信息公开平台，支持多数据源并行爬取
#
# ==============================================================================

"""
爬虫核心模块 - 政策爬取主逻辑
适配自然资源部政府信息公开平台
"""

import os
import json
import time
import logging
import hashlib
import threading
from typing import Dict, List, Optional, Callable, Any
from datetime import datetime, timezone
from urllib.parse import urljoin
from bs4 import BeautifulSoup

from .config import Config
from .api_client import APIClient
from .converter import DocumentConverter
from .models import Policy, CrawlProgress
from .mnr_spider import MNRSpider

# 使用模块级logger
logger = logging.getLogger(__name__)

# 文件编号锁（类级别，所有实例共享）
_file_number_lock = threading.Lock()

# 文件编号计数器（全局变量，线程安全）
_markdown_counter = 0
_file_counter = 0

# 自然资源部分类配置
MNR_CATEGORIES = {
    "自然资源调查监测": {"code": "1318", "name": "自然资源调查监测"},
    "自然资源确权登记": {"code": "1319", "name": "自然资源确权登记"},
    "自然资源合理开发利用": {"code": "1320", "name": "自然资源合理开发利用"},
    "自然资源有偿使用": {"code": "1321", "name": "自然资源有偿使用"},
    "国土空间规划": {"code": "1322", "name": "国土空间规划"},
    "国土空间用途管制": {"code": "1663", "name": "国土空间用途管制"},
    "国土空间生态修复": {"code": "1324", "name": "国土空间生态修复"},
    "耕地保护": {"code": "1325", "name": "耕地保护"},
    "地质勘查": {"code": "1326", "name": "地质勘查"},
    "矿产勘查": {"code": "1327", "name": "矿产勘查"},
    "矿产保护": {"code": "1328", "name": "矿产保护"},
    "矿产开发": {"code": "1329", "name": "矿产开发"},
    "地质环境保护": {"code": "1330", "name": "地质环境保护"},
    "海洋资源": {"code": "1331", "name": "海洋资源"},
    "测绘地理信息": {"code": "1332", "name": "测绘地理信息"},
    "地质灾害防治": {"code": "1334", "name": "地质灾害防治"},
    "地质公园": {"code": "1335", "name": "地质公园"},
    "地质遗迹保护": {"code": "1336", "name": "地质遗迹保护"},
    "矿业权评估": {"code": "1338", "name": "矿业权评估"},
    "机构建设": {"code": "1339", "name": "机构建设"},
    "综合管理": {"code": "1340", "name": "综合管理"},
    "其他": {"code": "1341", "name": "其他"},
}


class PolicyCrawler:
    """政策爬虫核心类 - 适配自然资源部API"""

    def __init__(self, config: Config, progress_callback: Optional[Callable] = None):
        """初始化爬虫

        Args:
            config: 配置对象
            progress_callback: 进度回调函数 (progress: CrawlProgress) -> None
        """
        self.config = config
        self.api_client = APIClient(config)
        self.converter = DocumentConverter()
        self.progress_callback = progress_callback
        self.stop_requested = False  # 停止标志
        self.progress = CrawlProgress()

        # 初始化 MNR 爬虫（使用新的核心实现，用于默认数据源）
        # 注意：在多数据源模式下，会为每个数据源创建新的爬虫实例
        self.mnr_spider = MNRSpider(config, self.api_client)
        self.base_url = self.mnr_spider.base_url

        # 创建输出目录
        self._create_output_dirs()

    def _create_output_dirs(self):
        """创建输出目录"""
        output_dir = self.config.output_dir
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(f"{output_dir}/json", exist_ok=True)
        os.makedirs(f"{output_dir}/files", exist_ok=True)
        os.makedirs(f"{output_dir}/markdown", exist_ok=True)
        os.makedirs(f"{output_dir}/docx", exist_ok=True)

    def request_stop(self):
        """请求停止爬取"""
        self.stop_requested = True
        logger.info("[停止] 收到停止请求，正在停止...")

    def _update_progress(self, **kwargs):
        """更新进度并触发回调"""
        for key, value in kwargs.items():
            setattr(self.progress, key, value)

        if self.progress_callback:
            # 生成详细的进度消息字符串
            current_stage = self.progress.get_current_stage()
            message_parts = []

            # 添加阶段信息
            if current_stage:
                message_parts.append(f"[{current_stage.description}]")

            # 添加当前处理信息
            if self.progress.current_policy_title:
                message_parts.append(f"正在处理: {self.progress.current_policy_title}")
            elif current_stage and current_stage.message:
                message_parts.append(current_stage.message)

            # 添加总体进度
            if self.progress.total_count > 0:
                completed = self.progress.completed_count + self.progress.failed_count
                message_parts.append(
                    f"总体进度: {completed}/{self.progress.total_count}"
                )

                # 添加阶段进度
                if current_stage and current_stage.total_count > 0:
                    stage_completed = (
                        current_stage.completed_count + current_stage.failed_count
                    )
                    message_parts.append(
                        f"阶段进度: {stage_completed}/{current_stage.total_count}"
                    )

            message = " | ".join(message_parts) if message_parts else "初始化中..."

            # 调用回调函数，传递字符串消息
            try:
                self.progress_callback(message)
            except Exception:
                # 如果回调函数不接受字符串，尝试传递对象（向后兼容）
                try:
                    self.progress_callback(self.progress)
                except Exception:
                    # 静默失败，避免中断爬虫流程
                    pass

    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """解析日期字符串为datetime对象"""
        if not date_str:
            return None
        for fmt in ("%Y年%m月%d日", "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(date_str, fmt)
            except Exception:
                continue
        return None

    def _parse_json_results(
        self, data: Dict, callback: Optional[Callable] = None
    ) -> List[Policy]:
        """解析JSON格式的搜索结果"""
        policies = []

        try:
            # 根据实际返回的JSON结构解析
            if "results" in data:
                items = data["results"]
            elif "data" in data:
                items = data["data"]
            elif isinstance(data, list):
                items = data
            else:
                items = []

            for item in items:
                # 尝试获取内容（如果API返回）
                content = item.get("content", "").strip()
                # 如果API没有返回内容，尝试从其他字段获取
                if not content:
                    content = (
                        item.get("summary", "").strip()
                        or item.get("abstract", "").strip()
                    )

                # 解析并格式化日期
                raw_date = item.get("pubdate", item.get("publishdate", ""))
                pub_date = ""
                if raw_date:
                    parsed_date = self._parse_date(raw_date)
                    if parsed_date:
                        pub_date = parsed_date.strftime("%Y-%m-%d")
                    else:
                        # 如果解析失败，尝试直接使用原始值（可能已经是标准格式）
                        pub_date = raw_date.strip()

                policy = Policy(
                    title=item.get("title", "") or "",
                    pub_date=pub_date or "",
                    doc_number=item.get("filenum", "") or "",
                    source=item.get("url", "") or "",
                    link=item.get("url", "") or "",
                    url=item.get("url", "") or "",
                    content=content or "",
                    category=item.get("category", "") or "",
                    validity=item.get("status", "") or "",
                    effective_date=item.get("effectivedate", "") or "",
                    crawl_time=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
                )
                policies.append(policy)

        except Exception as e:
            if callback:
                callback(f"解析JSON结果失败: {e}")

        return policies

    def _parse_html_results(
        self,
        soup: BeautifulSoup,
        callback: Optional[Callable] = None,
        category_name: str = "",
        max_policies: Optional[int] = None,
        data_source: Optional[Dict[str, Any]] = None,
    ) -> List[Policy]:
        """解析HTML格式的搜索结果

        Args:
            soup: BeautifulSoup对象
            callback: 进度回调函数
            category_name: 分类名称
            max_policies: 最大解析数量（None表示不限制）
        """
        policies = []

        try:
            # 查找政策列表 - 适配多个网站结构
            # f.mnr.gov.cn的结构：每个政策是一个独立的表格（5行）
            # 尝试查找所有表格
            tables = soup.find_all("table")

            if not tables:
                if callback:
                    callback("未找到任何表格")
                return policies

            # 对于f.mnr.gov.cn，每个政策是一个独立的表格
            # 检查表格结构：如果表格只有几行（通常是5行），且第一行包含"标题"，则是一个政策表格
            policy_tables = []
            for table in tables:
                rows = table.find_all("tr")
                # 政策表格通常有2-10行，且第一行第一列包含"标题"
                if len(rows) >= 2 and len(rows) <= 10:
                    # 检查第一行第一列是否包含"标题"等关键词
                    first_row = rows[0]
                    first_row_cells = first_row.find_all(["td", "th"])
                    if len(first_row_cells) >= 1:
                        first_cell_text = first_row_cells[0].get_text()
                        # 直接检查原始文本中是否包含"标"和"题"字符，或"名称"（f.mnr.gov.cn使用"名称"）
                        # 不进行lower，因为中文字符不需要
                        has_biao = "标" in first_cell_text
                        has_ti = "题" in first_cell_text
                        has_title = "标题" in first_cell_text
                        has_name = (
                            "名称" in first_cell_text
                            or "名称"
                            in first_cell_text.replace("\xa0", "").replace("\u00a0", "")
                        )

                        # 也检查清理后的文本
                        first_cell_clean = (
                            first_cell_text.replace("\xa0", "")
                            .replace("\u00a0", "")
                            .replace(" ", "")
                            .strip()
                        )
                        has_biao_clean = "标" in first_cell_clean
                        has_ti_clean = "题" in first_cell_clean
                        has_title_clean = "标题" in first_cell_clean
                        has_name_clean = "名称" in first_cell_clean

                        if (
                            has_title
                            or has_title_clean
                            or (has_biao and has_ti)
                            or (has_biao_clean and has_ti_clean)
                            or has_name
                            or has_name_clean
                        ):
                            policy_tables.append(table)
                            continue

                    # 如果第一列检查失败，尝试检查第一行的所有文本
                    first_row_text = first_row.get_text()
                    has_biao_row = "标" in first_row_text
                    has_ti_row = "题" in first_row_text
                    if ("标题" in first_row_text) or (has_biao_row and has_ti_row):
                        policy_tables.append(table)

            if not policy_tables:
                if callback:
                    callback(
                        f"未找到政策表格（共检查了 {len(tables)} 个表格），尝试查找其他结构..."
                    )
                    # 输出调试信息：检查前几个表格的结构
                    for i, table in enumerate(tables[:5], 1):
                        rows = table.find_all("tr")
                        if rows:
                            first_row = rows[0]
                            first_cells = first_row.find_all(["td", "th"])
                            if first_cells:
                                first_cell_text = first_cells[0].get_text()
                                callback(
                                    f"  表格{i}: {len(rows)}行, 第一列: {repr(first_cell_text[:30])}"
                                )

                # 如果找不到表格，尝试更宽松的条件：接受所有2-10行的表格
                if callback:
                    callback("尝试使用更宽松的条件识别政策表格...")
                for table in tables:
                    rows = table.find_all("tr")
                    if len(rows) >= 2 and len(rows) <= 10:
                        # 检查表格中是否有链接（政策表格通常包含链接）
                        links = table.find_all("a", href=True)
                        if links:
                            # 检查链接是否指向政策详情页
                            for link in links:
                                href = link.get("href", "")
                                if (
                                    href
                                    and not href.startswith("javascript")
                                    and (
                                        "html" in href
                                        or "detail" in href
                                        or "view" in href
                                    )
                                ):
                                    if table not in policy_tables:
                                        policy_tables.append(table)
                                    break

                if not policy_tables:
                    # 如果还是找不到，尝试查找列表结构（ul/ol/div）
                    list_containers = soup.find_all(
                        ["ul", "ol", "div"],
                        class_=lambda x: x
                        and ("list" in str(x).lower() or "item" in str(x).lower()),
                    )
                    if list_containers:
                        if callback:
                            callback(f"找到 {len(list_containers)} 个可能的列表容器")
                        # 这里可以添加列表解析逻辑，但先返回空列表
                        return policies
                    return policies

            if callback:
                callback(f"找到 {len(policy_tables)} 个政策表格")

            total_tables = len(policy_tables)

            # 如果设置了最大数量，只处理前N条有效政策
            parsed_count = 0

            # 遍历每个政策表格
            for table in policy_tables:
                # 如果已达到最大数量，停止解析
                if max_policies is not None and parsed_count >= max_policies:
                    break

                try:
                    # 解析政策表格 - f.mnr.gov.cn的结构
                    # 每个表格包含一个政策的详细信息，格式通常是：
                    # 第1行: [标签, 标题] - 标题行
                    # 第2行: [标签, 发文字号]
                    # 第3行: [标签, 发布日期]
                    # 第4行: [标签, 实施日期] (可选)
                    # 第5行: [标签, 其他信息] (可选)

                    rows = table.find_all("tr")
                    if len(rows) < 2:
                        continue

                    # 初始化变量
                    title = ""
                    doc_number = ""
                    pub_date = ""
                    detail_url = ""
                    level = "自然资源部"  # 默认值
                    validity = ""  # 默认值

                    # 遍历表格的每一行
                    for row in rows:
                        cells = row.find_all(["td", "th"])
                        if len(cells) < 2:
                            continue

                        label_raw = cells[0].get_text()
                        label = (
                            label_raw.replace("\xa0", "").replace(" ", "").strip()
                        )  # 清理特殊字符
                        value = cells[1].get_text(strip=True)

                        # 根据标签提取信息（使用清理后的标签）
                        # 支持"标题"和"名称"两种标签
                        if (
                            "标题" in label
                            or "名称" in label
                            or "title" in label.lower()
                            or ("标" in label and "题" in label)
                            or ("名" in label and "称" in label)
                        ):
                            title = value
                            # 查找标题中的链接
                            link = cells[1].find("a", href=True)
                            if link:
                                detail_url = link.get("href", "")
                                # 如果标题为空，使用链接文本
                                if not title:
                                    title = link.get_text(strip=True)
                        elif (
                            "发文字号" in label
                            or "文号" in label
                            or ("发" in label and "号" in label)
                        ):
                            doc_number = value
                        elif (
                            "成文时间" in label
                            or "生成日期" in label
                            or "发布日期" in label
                            or "公布日期" in label
                        ) and not ("效力" in label or "级别" in label):
                            # 验证值是否是日期格式（包含年月日）
                            if (
                                any(keyword in value for keyword in ["年", "月", "日"])
                                or len(value) >= 8
                            ):
                                pub_date = value
                        elif (
                            "实施日期" in label
                            or "生效日期" in label
                            or (
                                "日期" in label and ("实施" in label or "生效" in label)
                            )
                        ):
                            # 如果没有发布日期，使用实施日期
                            if not pub_date:
                                pub_date = value
                        elif "发布机构" in label or "机构" in label:
                            # 提取发布机构
                            if value and value.strip():
                                level = value
                        elif "效力级别" in label or "级别" in label:
                            # 提取效力级别
                            if value and value.strip():
                                validity = value

                    # 如果标题为空，尝试从表格中查找所有链接
                    if not title:
                        links = table.find_all("a", href=True)
                        for link in links:
                            link_text = link.get_text(strip=True)
                            link_href = link.get("href", "")
                            # 跳过javascript链接
                            if link_href and not link_href.startswith("javascript"):
                                if len(link_text) > 5:  # 标题应该有一定长度
                                    title = link_text
                                    detail_url = link_href
                                    break

                    if not title:
                        continue

                    # 构建完整链接
                    if detail_url and not detail_url.startswith("http"):
                        detail_url = urljoin(self.base_url, detail_url)

                    # 解析并格式化日期
                    pub_date_formatted = ""
                    if pub_date:
                        parsed_date = self._parse_date(pub_date)
                        if parsed_date:
                            pub_date_formatted = parsed_date.strftime("%Y-%m-%d")
                        else:
                            # 如果解析失败，尝试直接使用原始值（可能已经是标准格式）
                            pub_date_formatted = pub_date.strip()

                    policy = Policy(
                        title=title,
                        pub_date=pub_date_formatted,
                        doc_number=doc_number,
                        source=detail_url,
                        link=detail_url,
                        url=detail_url,
                        content="",  # 初始为空，会在crawl_policies中填充
                        category=category_name,
                        level=level,
                        validity=validity,
                        crawl_time=datetime.now(timezone.utc).strftime(
                            "%Y-%m-%d %H:%M:%S"
                        ),
                    )
                    # 保存数据源信息到policy对象（通过添加自定义属性）
                    policy._data_source = data_source

                    policies.append(policy)
                    parsed_count += 1  # 成功解析一条政策

                except Exception as e:
                    if callback:
                        callback(f"解析政策项失败: {e}")
                    continue

            # 统计信息
            if callback:
                callback(
                    f"成功解析 {len(policies)} 条政策（从 {total_tables} 个表格中）"
                )

        except Exception as e:
            if callback:
                callback(f"解析HTML结果失败: {e}")

        return policies

    def search_all_policies(
        self,
        keywords: List[str] = None,
        start_date: str = None,
        end_date: str = None,
        callback: Optional[Callable] = None,
        limit_pages: Optional[int] = None,
    ) -> List[Policy]:
        """搜索所有政策（支持多数据源顺序执行）

        Args:
            keywords: 关键词列表
            start_date: 起始日期 yyyy-MM-dd
            end_date: 结束日期 yyyy-MM-dd
            callback: 进度回调函数
            limit_pages: 限制最大页数（用于测试模式，None表示不限制）

        Returns:
            政策列表
        """
        if keywords is None:
            keywords = []

        # 获取启用的数据源列表
        data_sources = self.config.get("data_sources", [])
        if not data_sources:
            # 如果没有配置数据源列表，使用默认配置（向后兼容）
            data_sources = [
                {
                    "name": "政府信息公开平台",
                    "base_url": self.config.get("base_url", "https://gi.mnr.gov.cn/"),
                    "search_api": self.config.get(
                        "search_api", "https://search.mnr.gov.cn/was5/web/search"
                    ),
                    "channel_id": self.config.get("channel_id", "216640"),
                    "enabled": True,
                }
            ]

        # 过滤出启用的数据源
        enabled_sources = [ds for ds in data_sources if ds.get("enabled", True)]

        if not enabled_sources:
            if callback:
                callback("错误：没有启用的数据源，请至少选择一个数据源")
            return []

        if callback:
            if keywords:
                callback(f"开始搜索政策，关键词: {', '.join(keywords)}")
            else:
                callback("开始全量爬取：无关键词限制，将爬取所有政策")

            if start_date or end_date:
                date_info = f"时间范围: {start_date or '不限'} 至 {end_date or '不限'}"
                callback(date_info)
            else:
                callback("时间范围: 无限制（全量爬取）")

            callback(
                f"启用数据源: {', '.join([ds.get('name', '未知') for ds in enabled_sources])}"
            )
            if len(enabled_sources) > 1:
                callback(f"将按顺序执行 {len(enabled_sources)} 个数据源")

        all_policies = []
        seen_ids = set()  # 用于跨数据源去重

        def stop_check():
            return self.stop_requested

        # 按顺序遍历每个启用的数据源
        for idx, data_source in enumerate(enabled_sources, 1):
            if self.stop_requested:
                if callback:
                    callback("停止搜索")
                break

            source_name = data_source.get("name", f"数据源{idx}")
            if callback:
                callback(f"\n{'='*60}")
                callback(f"开始爬取数据源 {idx}/{len(enabled_sources)}: {source_name}")
                callback(f"{'='*60}")

            # 为当前数据源创建临时配置
            temp_config = Config()
            temp_config.config = self.config.config.copy()
            # 确保数据源明确标记为启用
            data_source_copy = data_source.copy()
            data_source_copy["enabled"] = True
            temp_config.config["data_sources"] = [data_source_copy]
            temp_config.config["base_url"] = data_source.get(
                "base_url", "https://gi.mnr.gov.cn/"
            )
            temp_config.config["search_api"] = data_source.get(
                "search_api", "https://search.mnr.gov.cn/was5/web/search"
            )
            temp_config.config["channel_id"] = data_source.get("channel_id", "216640")

            # 为当前数据源创建MNR爬虫（会自动选择对应的HTML解析器）
            # 注意：MNRSpider已在文件顶部导入，这里重新导入会导致变量遮蔽
            temp_api_client = APIClient(temp_config)
            mnr_spider = MNRSpider(temp_config, temp_api_client)

            if callback:
                parser_type = type(mnr_spider.html_parser).__name__
                callback(
                    f"使用解析器: {parser_type} (适配 {data_source.get('base_url', '')})"
                )

            # 如果设置了 limit_pages，临时修改 max_pages
            original_max_pages = mnr_spider.max_pages
            if limit_pages is not None:
                mnr_spider.max_pages = limit_pages

            try:
                policies = mnr_spider.crawl_policies(
                    keywords=keywords,
                    callback=callback,
                    start_date=start_date,
                    end_date=end_date,
                    category=None,  # 搜索全部分类
                    stop_callback=stop_check,
                    policy_callback=None,
                )

                # 去重并添加到总列表
                source_policy_count = 0
                for policy in policies:
                    policy_id = policy.id
                    if policy_id not in seen_ids:
                        seen_ids.add(policy_id)
                        all_policies.append(policy)
                        source_policy_count += 1

                if callback:
                    callback(
                        f"数据源 {source_name} 完成，获取 {len(policies)} 条政策（新增 {source_policy_count} 条，去重后）"
                    )

            except Exception as e:
                # 捕获异常，记录错误但继续处理下一个数据源
                error_msg = f"数据源 {source_name} 爬取失败: {str(e)}"
                if callback:
                    callback(f"[错误] {error_msg}")
                logger.error(f"数据源 {source_name} 爬取异常: {e}", exc_info=True)

                # 记录数据源级别的失败（如果获取到政策但爬取失败）
                # 使用logger记录失败信息
                logger.error(f"数据源 {source_name} 爬取失败: {error_msg}")
                # 继续处理下一个数据源，不中断整个流程

            finally:
                # 恢复原始 max_pages
                mnr_spider.max_pages = original_max_pages
                # 关闭临时API客户端
                if hasattr(temp_api_client, "close"):
                    try:
                        temp_api_client.close()
                    except Exception:
                        pass  # 忽略关闭时的异常

        if callback:
            callback(f"\n所有数据源爬取完成，总计 {len(all_policies)} 条政策（已去重）")

        return all_policies

    def crawl_single_policy(
        self, policy: Policy, callback: Optional[Callable] = None, retry_count: int = 0
    ) -> Optional[Policy]:
        """爬取单个政策（支持自动重试）

        Args:
            policy: 政策对象
            callback: 进度回调函数
            retry_count: 当前重试次数（内部使用）

        Returns:
            爬取成功的政策对象，失败时返回None
        """
        self._update_progress(
            current_policy_id=policy.id, current_policy_title=policy.title
        )

        if retry_count > 0:
            if callback:
                callback(f"\n[重试 {retry_count}] 爬取政策: {policy.title}")
            logger.info(f"[重试 {retry_count}] 爬取政策: {policy.title}")
        else:
            if callback:
                callback(f"\n爬取政策: {policy.title}")
            logger.info(f"爬取政策: {policy.title}")

        max_policy_retries = self.config.get("max_policy_retries", 0)  # 默认不重试

        try:
            # 1. 获取详情页内容和附件（如果还没有）
            attachments = []
            if not policy.content and policy.link:
                if callback:
                    callback("获取详情页内容和附件...")
                logger.debug(f"获取政策详情页: {policy.title[:50]}...")
                # 获取数据源信息（如果policy对象有保存）
                data_source = getattr(policy, "_data_source", None)
                detail_result = self.api_client.get_policy_detail(
                    policy.link, data_source
                )

                if not detail_result:
                    raise Exception("获取详情页失败：无响应")

                policy.content = detail_result.get("content", "")
                attachments = detail_result.get("attachments", [])

                # 更新元信息（如果详情页有更完整的信息）
                metadata = detail_result.get("metadata", {})
                if metadata:
                    # 更新发布日期（如果详情页有且列表页没有，或列表页的值不是日期格式）
                    if metadata.get("pub_date"):
                        # 验证是否是日期格式
                        is_date_format = (
                            any(
                                keyword in metadata["pub_date"]
                                for keyword in ["年", "月", "日"]
                            )
                            or len(metadata["pub_date"]) >= 8
                        )
                        if is_date_format:
                            # 如果列表页没有，或者列表页的值不是日期格式，则更新
                            if not policy.pub_date or not any(
                                keyword in policy.pub_date
                                for keyword in ["年", "月", "日"]
                            ):
                                parsed_date = self._parse_date(metadata["pub_date"])
                                if parsed_date:
                                    policy.pub_date = parsed_date.strftime("%Y-%m-%d")
                                else:
                                    policy.pub_date = metadata["pub_date"].strip()
                    # 更新发布机构（如果详情页有）
                    if metadata.get("publisher"):
                        policy.publisher = metadata["publisher"]
                    # 更新效力级别（level字段，如果详情页有）
                    if metadata.get("level"):
                        policy.level = metadata["level"]
                    # 更新有效性（validity字段，如果详情页有）
                    if metadata.get("validity"):
                        policy.validity = metadata["validity"]
                    # 更新生效日期（如果详情页有）
                    if metadata.get("effective_date"):
                        parsed_date = self._parse_date(metadata["effective_date"])
                        if parsed_date:
                            policy.effective_date = parsed_date.strftime("%Y-%m-%d")
                        else:
                            policy.effective_date = metadata["effective_date"].strip()
                    # 更新分类（如果详情页有）
                    if metadata.get("category"):
                        policy.category = metadata["category"]
                    # 更新文号（如果详情页有且列表页没有）
                    if metadata.get("doc_number") and not policy.doc_number:
                        policy.doc_number = metadata["doc_number"]

                # 后备逻辑：根据文号推断发布机构和效力级别
                if policy.doc_number and not policy.publisher:
                    if policy.doc_number.startswith("自然资发"):
                        policy.publisher = "自然资源部"
                        if not policy.level:
                            policy.level = "自然资源部"
                        logger.debug(
                            f"根据文号推断发布机构: {policy.doc_number} -> 自然资源部"
                        )
                    elif policy.doc_number.startswith(
                        "国土资发"
                    ) or policy.doc_number.startswith("国土调查办发"):
                        policy.publisher = "国土资源部"
                        if not policy.level:
                            policy.level = "国土资源部"
                        logger.debug(
                            f"根据文号推断发布机构: {policy.doc_number} -> 国土资源部"
                        )

            # 2. 保存JSON数据
            if self.config.get("save_json", True):
                self._save_json(policy)

            # 3. 获取文件编号（markdown 和 files 文件夹各自独立递增）
            markdown_number = self._get_next_markdown_number()
            file_number = self._get_next_file_number()

            # 4. 下载附件（如果启用）
            if self.config.get("save_files", True) and attachments:
                self._download_attachments(policy, attachments, file_number, callback)

            # 5. 生成RAG Markdown
            save_markdown = self.config.get("save_markdown", True)
            logger.info(f"爬虫配置: save_markdown={save_markdown}")
            if save_markdown:
                logger.info("开始生成Markdown文件...")
                self._generate_rag_markdown(policy, markdown_number)
                logger.info(f"Markdown文件生成完成: {policy.markdown_path}")
            else:
                logger.warning("Markdown文件生成已禁用")

            # 6. 生成DOCX格式（如果启用）
            save_docx = self.config.get("save_docx", True)
            logger.info(f"爬虫配置: save_docx={save_docx}")
            if save_docx:
                logger.info("开始生成DOCX文件...")
                self._generate_docx(policy, markdown_number, callback)
                logger.info("DOCX文件生成完成")
            else:
                logger.warning("DOCX文件生成已禁用")

            if callback:
                callback("   ✓ 政策爬取完成")
            logger.debug("政策爬取完成")
            return policy

        except Exception as e:
            # 记录失败信息到专门的失败日志
            error_msg = str(e)

            # 检查是否需要重试
            max_policy_retries = self.config.get("max_policy_retries", 0)
            if retry_count < max_policy_retries:
                # 自动重试
                retry_delay = self.config.get("policy_retry_delay", 5)
                if callback:
                    callback(f"   ✗ 政策爬取失败: {error_msg}")
                    callback(
                        f"   [自动重试] 等待 {retry_delay} 秒后重试 ({retry_count + 1}/{max_policy_retries})..."
                    )
                logger.warning(
                    f"政策爬取失败，将自动重试 ({retry_count + 1}/{max_policy_retries}): {policy.title} - {error_msg}"
                )

                time.sleep(retry_delay)
                retry_result = self.crawl_single_policy(
                    policy, callback, retry_count + 1
                )
                return retry_result
            else:
                # 重试次数已用完，记录失败
                logger.error(
                    f"政策爬取失败 (已重试 {retry_count} 次): {policy.title} - {error_msg}"
                )

                if callback:
                    callback(
                        f"   ✗ 政策爬取失败: {error_msg} (已重试 {retry_count} 次)"
                    )
                logger.error(
                    f"政策爬取失败 (已重试 {retry_count} 次): {policy.title} - {error_msg}",
                    exc_info=True,
                )
                return None

    def _save_json(self, policy: Policy):
        """保存JSON数据"""
        # 生成安全的文件名：如果ID过长或包含特殊字符，使用hash摘要
        original_id = str(policy.id)
        safe_id = (
            original_id.replace("|", "_")
            .replace("/", "_")
            .replace("\\", "_")
            .replace(":", "_")
        )

        # 如果ID过长或包含不安全字符，使用hash
        if len(safe_id) > 100 or not all(
            c.isalnum() or c in ("_", "-", ".") for c in safe_id
        ):
            # 使用SHA256的前16个字符作为文件名（足够唯一且简短）
            id_hash = hashlib.sha256(original_id.encode("utf-8")).hexdigest()[:16]
            policy_id = f"{safe_id[:20]}_{id_hash}" if len(safe_id) > 0 else id_hash
        else:
            policy_id = safe_id[:100]  # 限制长度

        filepath = f"{self.config.output_dir}/json/policy_{policy_id}.json"

        try:
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(policy.to_dict(), f, ensure_ascii=False, indent=2)
            logger.debug(f"JSON已保存: {filepath}")
        except Exception as e:
            logger.error(f"JSON保存失败: {e}", exc_info=True)

    def _generate_rag_markdown(
        self, policy: Policy, file_number: int, callback: Optional[Callable] = None
    ):
        """生成RAG格式的Markdown文件"""
        try:
            md_lines = []

            # YAML Front Matter
            md_lines.append("---")
            md_lines.append(f'title: "{policy.title}"')
            md_lines.append(f'level: "{policy.level}"')
            md_lines.append(f'category: "{policy.category}"')
            md_lines.append(f'pub_date: "{policy.pub_date}"')
            md_lines.append(f'doc_number: "{policy.doc_number}"')
            if policy.effective_date:
                md_lines.append(f'effective_date: "{policy.effective_date}"')
            if policy.validity:
                md_lines.append(f'validity: "{policy.validity}"')
            md_lines.append(f'source_url: "{policy.source}"')
            md_lines.append(f'crawl_time: "{policy.crawl_time}"')
            md_lines.append("---")
            md_lines.append("")

            md_lines.append(f"# {policy.title}")
            md_lines.append("")

            md_lines.append("## 基本信息")
            md_lines.append("")
            md_lines.append(f"- **发布机构**: {policy.level}")
            md_lines.append(f"- **发布日期**: {policy.pub_date}")
            if policy.doc_number:
                md_lines.append(f"- **发文字号**: {policy.doc_number}")
            if policy.effective_date:
                md_lines.append(f"- **生效日期**: {policy.effective_date}")
            if policy.validity:
                md_lines.append(f"- **有效性**: {policy.validity}")
            if policy.category:
                md_lines.append(f"- **分类**: {policy.category}")
            md_lines.append(f"- **来源链接**: [查看原文]({policy.source})")
            md_lines.append("")

            md_lines.append("---")
            md_lines.append("")
            md_lines.append("## 正文内容")
            md_lines.append("")
            if policy.content:
                md_lines.append(policy.content)
            else:
                md_lines.append("> **注意**: 该政策的正文内容无法自动获取。")
                md_lines.append("> ")
                md_lines.append("> 请访问[来源链接](#基本信息)查看完整文档内容。")

            safe_title = "".join(
                c for c in policy.title if c.isalnum() or c in (" ", "-", "_")
            ).strip()
            if not safe_title:
                safe_title = f"政策_{policy.id[:8]}"

            md_filename = f"{file_number:04d}_{safe_title}.md"
            md_filepath = f"{self.config.output_dir}/markdown/{md_filename}"

            with open(md_filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(md_lines))

            # 保存文件路径到policy对象，以便后续保存到数据库
            policy.markdown_path = md_filepath
            if callback:
                callback(f"  [OK] Markdown已保存: {md_filename}")
            logger.debug(f"Markdown已保存: {md_filepath}")

        except Exception as e:
            logger.error(f"Markdown生成失败: {e}", exc_info=True)

    def _generate_docx(
        self, policy: Policy, file_number: int, callback: Optional[Callable] = None
    ):
        """生成DOCX格式文件

        Args:
            policy: 政策对象
            file_number: 文件编号
            callback: 进度回调函数
        """
        try:
            # 尝试导入 python-docx
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                if callback:
                    callback("  [X] python-docx未安装，无法生成DOCX")
                logger.warning("python-docx未安装，无法生成DOCX")
                return

            # 创建文档
            doc = Document()

            # 设置文档标题
            title = doc.add_heading(policy.title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加基本信息
            doc.add_heading("基本信息", level=2)
            info_para = doc.add_paragraph()
            info_para.add_run("发布机构: ").bold = True
            info_para.add_run(policy.level)
            info_para.add_run("\n发布日期: ").bold = True
            info_para.add_run(policy.pub_date)

            if policy.doc_number:
                info_para.add_run("\n发文字号: ").bold = True
                info_para.add_run(policy.doc_number)

            if policy.effective_date:
                info_para.add_run("\n生效日期: ").bold = True
                info_para.add_run(policy.effective_date)

            if policy.validity:
                info_para.add_run("\n有效性: ").bold = True
                info_para.add_run(policy.validity)

            if policy.category:
                info_para.add_run("\n分类: ").bold = True
                info_para.add_run(policy.category)

            info_para.add_run("\n来源链接: ").bold = True
            info_para.add_run(policy.source)

            # 添加正文内容
            doc.add_heading("正文内容", level=2)

            if policy.content:
                # 将正文内容按段落分割并添加
                content_lines = policy.content.split("\n")
                for line in content_lines:
                    line = line.strip()
                    if line:
                        # 检查是否是标题（以#开头）
                        if line.startswith("#"):
                            # 计算标题级别
                            level = len(line) - len(line.lstrip("#"))
                            heading_text = line.lstrip("#").strip()
                            if heading_text:
                                doc.add_heading(heading_text, level=min(level, 9))
                        else:
                            doc.add_paragraph(line)
            else:
                doc.add_paragraph("该政策的正文内容无法自动获取。")
                doc.add_paragraph(f"请访问来源链接查看完整文档内容: {policy.source}")

            # 保存文件
            safe_title = "".join(
                c for c in policy.title if c.isalnum() or c in (" ", "-", "_")
            ).strip()
            if not safe_title:
                safe_title = f"政策_{policy.id[:8]}"

            docx_filename = f"{file_number:04d}_{safe_title}.docx"
            docx_filepath = f"{self.config.output_dir}/docx/{docx_filename}"

            doc.save(docx_filepath)

            # 保存文件路径到policy对象，以便后续保存到数据库
            policy.docx_path = docx_filepath
            logger.debug(f"DOCX已保存: {docx_filepath}")

        except Exception as e:
            if callback:
                callback(f"  [X] DOCX生成失败: {e}")
            logger.error(f"DOCX生成失败: {e}", exc_info=True)

    def _get_next_markdown_number(self) -> int:
        """获取下一个 Markdown 文件编号（线程安全）"""
        global _markdown_counter

        with _file_number_lock:
            # 如果计数器为0，从文件系统初始化
            if _markdown_counter == 0:
                markdown_dir = f"{self.config.output_dir}/markdown"
                if os.path.exists(markdown_dir):
                    existing_files = [
                        f for f in os.listdir(markdown_dir) if f.endswith(".md")
                    ]
                    if existing_files:
                        numbers = []
                        for filename in existing_files:
                            parts = filename.split("_", 1)
                            if parts and len(parts) >= 2 and parts[0].isdigit():
                                numbers.append(int(parts[0]))
                        if numbers:
                            _markdown_counter = max(numbers)

            # 递增计数器并返回新值
            _markdown_counter += 1
            return _markdown_counter

    def _get_next_file_number(self) -> int:
        """获取下一个附件文件编号（线程安全）"""
        global _file_counter

        with _file_number_lock:
            # 如果计数器为0，从文件系统初始化
            if _file_counter == 0:
                files_dir = f"{self.config.output_dir}/files"
                if os.path.exists(files_dir):
                    existing_files = os.listdir(files_dir)
                    if existing_files:
                        numbers = []
                        for filename in existing_files:
                            parts = filename.split("_", 1)
                            if parts and len(parts) >= 2 and parts[0].isdigit():
                                numbers.append(int(parts[0]))
                        if numbers:
                            _file_counter = max(numbers)

            # 递增计数器并返回新值
            _file_counter += 1
            return _file_counter

    def _download_attachments(
        self,
        policy: Policy,
        attachments: List[Dict[str, str]],
        file_number: int,
        callback: Optional[Callable] = None,
    ):
        """下载附件

        Args:
            policy: 政策对象
            attachments: 附件列表，每个附件包含 {'url': str, 'name': str}
            file_number: 文件编号
            callback: 进度回调函数
        """
        if not attachments:
            return

        # 筛选需要下载的文件
        target_files = []

        if self.config.get("download_all_files", False):
            target_files = attachments
            if callback:
                callback(
                    f"\n[下载所有文件] 已启用，将下载所有 {len(attachments)} 个附件"
                )
        else:
            # 根据文件扩展名筛选
            for att in attachments:
                url = att.get("url", "")
                name = att.get("name", "")

                # 从URL或名称中提取扩展名
                file_ext = ""
                if url:
                    file_ext = os.path.splitext(url.split("?")[0])[1].lower().strip(".")
                if not file_ext and name:
                    file_ext = os.path.splitext(name)[1].lower().strip(".")

                should_download = False
                if file_ext == "docx" and self.config.get("download_docx", True):
                    should_download = True
                elif file_ext == "doc" and self.config.get("download_doc", True):
                    should_download = True
                elif file_ext == "pdf" and self.config.get("download_pdf", False):
                    should_download = True
                elif file_ext in ["zip", "tar", "rar", "7z", "gz"]:
                    # 压缩文件总是下载（包括.tar, .zip等）
                    should_download = True

                if should_download:
                    target_files.append(att)

            if target_files:
                if callback:
                    callback(
                        f"\n从 {len(attachments)} 个附件中筛选出 {len(target_files)} 个文件"
                    )

        if not target_files:
            if callback:
                callback("没有需要下载的附件")
            return

        # 下载文件
        safe_title = "".join(
            c for c in policy.title if c.isalnum() or c in (" ", "-", "_")
        ).strip()
        if not safe_title:
            safe_title = f"政策_{policy.id[:8]}"

        for i, att in enumerate(target_files, 1):
            url = att.get("url", "")
            name = att.get("name", "")

            if callback:
                callback(f"\n  [{i}/{len(target_files)}] 下载: {name or url}")

            # 确定保存文件名
            if name:
                # 清理文件名
                safe_name = "".join(
                    c for c in name if c.isalnum() or c in (" ", "-", "_", ".")
                ).strip()
                if not safe_name:
                    safe_name = f"附件_{i}"
                # 确保有扩展名
                if not os.path.splitext(safe_name)[1]:
                    # 从URL提取扩展名
                    ext = os.path.splitext(url.split("?")[0])[1]
                    if ext:
                        safe_name += ext
            else:
                # 从URL提取文件名
                safe_name = url.split("/")[-1].split("?")[0]
                if not safe_name or len(safe_name) < 3:
                    safe_name = f"附件_{i}"

            # 智能文件名生成：如果只有一个附件且名称与政策标题相似，使用简化名称
            # 使用模糊匹配：比较文件名主体（去除扩展名和常见后缀）
            def get_base_name(name):
                """提取文件名主体，去除扩展名和常见后缀"""
                base = os.path.splitext(name)[0]
                # 去除常见的后缀如"附件"、"下载"等
                suffixes = ["附件", "下载", "附件1", "文件"]
                for suffix in suffixes:
                    if base.endswith(suffix):
                        base = base[: -len(suffix)]
                return base.strip()

            policy_title_base = get_base_name(policy.title)
            safe_name_base = get_base_name(safe_name)

            # 如果只有一个附件且文件名主体与政策标题相似（允许部分匹配），使用简化名称
            if len(target_files) == 1:
                # 检查文件名是否包含标题的主要部分，或标题是否包含文件名
                if (
                    policy_title_base
                    and safe_name_base
                    and (
                        policy_title_base in safe_name_base
                        or safe_name_base in policy_title_base
                        or len(set(policy_title_base) & set(safe_name_base))
                        / max(len(policy_title_base), len(safe_name_base), 1)
                        > 0.6
                    )
                ):
                    ext = (
                        os.path.splitext(safe_name)[1]
                        or os.path.splitext(url.split("?")[0])[1]
                    )
                    save_filename = f"{file_number:04d}_{safe_title}{ext}"
                else:
                    save_filename = f"{file_number:04d}_{safe_title}_{safe_name}"
            else:
                save_filename = f"{file_number:04d}_{safe_title}_{safe_name}"

            save_path = f"{self.config.output_dir}/files/{save_filename}"

            # 下载文件
            if self.api_client.download_file(url, save_path):
                if callback:
                    callback(f"    [OK] 下载成功: {save_filename}")
                logger.debug(f"附件下载成功: {save_path}")
                # 保存附件路径到policy对象，以便后续保存到数据库
                if not hasattr(policy, "_attachment_paths"):
                    policy._attachment_paths = []
                policy._attachment_paths.append(
                    {
                        "url": url,
                        "name": name,
                        "storage_path": save_path,
                        "file_name": save_filename,
                    }
                )
            else:
                if callback:
                    callback(f"    [X] 下载失败: {name or url}")
                logger.warning(f"附件下载失败: {url}")

            # 下载间隔
            if i < len(target_files):
                time.sleep(0.3)

    def crawl_batch(
        self,
        keywords: List[str] = None,
        start_date: str = None,
        end_date: str = None,
        callback: Optional[Callable] = None,
    ) -> CrawlProgress:
        """批量爬取"""
        if not hasattr(self, "progress") or self.progress is None:
            self.progress = CrawlProgress()
            self.progress.start_time = datetime.now(timezone.utc)
        elif self.progress.start_time is None:
            self.progress.start_time = datetime.now(timezone.utc)
        else:
            original_start_time = self.progress.start_time
            self.progress = CrawlProgress()
            self.progress.start_time = original_start_time

        self._update_progress()

        # 设置搜索阶段
        self.progress.set_stage("search_policies", "搜索政策列表")
        self._update_progress()

        # 1. 搜索所有政策
        all_policies = self.search_all_policies(
            keywords, start_date, end_date, callback
        )

        if self.stop_requested:
            self.progress.end_time = datetime.now(timezone.utc)
            self._update_progress()
            return self.progress

        # 设置详情爬取阶段
        self.progress.set_stage(
            "crawl_details",
            f"爬取政策详情（共 {len(all_policies)} 条）",
            len(all_policies),
        )
        self._update_progress()

        if callback:
            callback(f"\n开始爬取政策详细内容，共 {len(all_policies)} 条政策")
        logger.info(f"开始爬取政策详细内容，共 {len(all_policies)} 条政策")

        # 2. 爬取每个政策的详细内容
        for i, policy in enumerate(all_policies, 1):
            if self.stop_requested:
                if callback:
                    callback("停止爬取政策")
                logger.info("[停止] 停止爬取政策")
                break

            if callback:
                callback(f"\n进度: [{i}/{len(all_policies)}]")
            logger.info(f"进度: [{i}/{len(all_policies)}]")

            self.progress.current_policy_id = policy.id
            self.progress.current_policy_title = policy.title
            self._update_progress()

            result_policy = self.crawl_single_policy(policy, callback)

            if result_policy:
                self.progress.update_stage_progress("crawl_details", completed=1)
                self.progress.completed_policies.append(policy.id)
                if callback:
                    callback(f"✓ {policy.title}")
            else:
                self.progress.update_stage_progress("crawl_details", failed=1)
                # 从失败日志中获取失败原因（如果已记录）
                failure_reason = "爬取失败"
                failed_policy_info = {
                    "id": policy.id,
                    "title": policy.title,
                    "link": policy.link or policy.source or policy.url,
                    "pub_date": policy.pub_date,
                    "doc_number": policy.doc_number,
                    "reason": failure_reason,
                }
                self.progress.failed_policies.append(failed_policy_info)
                if callback:
                    callback(f"✗ {policy.title} - {failure_reason}")

            self._update_progress()
            time.sleep(self.config.get("request_delay", 2))

        self.progress.end_time = datetime.now(timezone.utc)
        self._update_progress()

        if callback:
            callback("\n爬取完成")
        logger.info("爬取完成")
        logger.info(f"总计: {self.progress.total_count} 条")
        logger.info(f"成功: {self.progress.completed_count} 条")
        logger.info(f"失败: {self.progress.failed_count} 条")
        logger.info(f"成功率: {self.progress.success_rate:.2f}%")

        # 如果有失败的政策，提示查看失败日志
        if self.progress.failed_count > 0:
            from pathlib import Path

            log_dir = self.config.get("log_dir", "logs")
            log_dir_path = Path(log_dir)
            log_dir_path.mkdir(parents=True, exist_ok=True)
            failure_log_file = (
                log_dir_path
                / f"failures_{datetime.now(timezone.utc).strftime('%Y%m%d')}.log"
            )

            if callback:
                callback(f"\n失败的政策已记录到: {failure_log_file}")
            logger.info(f"失败的政策已记录到: {failure_log_file}")

        return self.progress

    def retry_failed_policies(
        self,
        failure_log_file: Optional[str] = None,
        callback: Optional[Callable] = None,
    ) -> CrawlProgress:
        """重试失败的政策（从失败日志文件中读取）

        Args:
            failure_log_file: 失败日志文件路径（如果为None，使用今天的失败日志）
            callback: 进度回调函数

        Returns:
            爬取进度
        """
        from pathlib import Path

        # 确定失败日志文件路径
        if failure_log_file is None:
            log_dir = self.config.get("log_dir", "logs")
            log_dir_path = Path(log_dir)
            failure_log_file = (
                log_dir_path
                / f"failures_{datetime.now(timezone.utc).strftime('%Y%m%d')}.log"
            )
        else:
            failure_log_file = Path(failure_log_file)

        if not failure_log_file.exists():
            if callback:
                callback(f"失败日志文件不存在: {failure_log_file}")
            logger.warning(f"失败日志文件不存在: {failure_log_file}")
            return CrawlProgress()

        if callback:
            callback(f"\n读取失败日志: {failure_log_file}")
        logger.info(f"读取失败日志: {failure_log_file}")

        # 解析失败日志文件
        failed_policies = []
        try:
            with open(failure_log_file, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or "|" not in line:
                        continue

                    # 解析日志格式: 时间 | 标题: xxx | 链接: xxx | 发布日期: xxx | 发文字号: xxx | 失败原因: xxx
                    # 或者: 时间 | 标题: xxx | 链接: xxx | 发布日期: | 发文字号: | 失败原因: xxx
                    parts = line.split("|")
                    if len(parts) < 6:
                        continue

                    # 提取信息
                    title = ""
                    link = ""
                    pub_date = ""
                    doc_number = ""
                    reason = ""

                    for part in parts[1:]:  # 跳过时间部分
                        part = part.strip()
                        if part.startswith("标题:"):
                            title = part.replace("标题:", "").strip()
                        elif part.startswith("链接:"):
                            link = part.replace("链接:", "").strip()
                        elif part.startswith("发布日期:"):
                            pub_date = part.replace("发布日期:", "").strip()
                        elif part.startswith("发文字号:"):
                            doc_number = part.replace("发文字号:", "").strip()
                        elif part.startswith("失败原因:"):
                            reason = part.replace("失败原因:", "").strip()

                    if title and link:
                        policy = Policy(
                            title=title,
                            pub_date=pub_date,
                            doc_number=doc_number,
                            source=link,
                            link=link,
                            url=link,
                            content="",
                            category="",
                            level="自然资源部",
                            validity="",
                            effective_date="",
                            crawl_time=datetime.now(timezone.utc).strftime(
                                "%Y-%m-%d %H:%M:%S"
                            ),
                        )
                        failed_policies.append((policy, reason))

        except Exception as e:
            if callback:
                callback(f"解析失败日志文件失败: {e}")
            logger.error(f"解析失败日志文件失败: {e}", exc_info=True)
            return CrawlProgress()

        if not failed_policies:
            if callback:
                callback("失败日志文件中没有找到失败的政策")
            logger.info("失败日志文件中没有找到失败的政策")
            return CrawlProgress()

        if callback:
            callback(f"\n找到 {len(failed_policies)} 条失败的政策，开始重试...")
        logger.info(f"找到 {len(failed_policies)} 条失败的政策，开始重试...")

        # 初始化进度
        if not hasattr(self, "progress") or self.progress is None:
            self.progress = CrawlProgress()
            self.progress.start_time = datetime.now(timezone.utc)
        else:
            original_start_time = self.progress.start_time
            self.progress = CrawlProgress()
            self.progress.start_time = original_start_time or datetime.now(timezone.utc)

        self.progress.total_count = len(failed_policies)
        self._update_progress()

        # 重试每个失败的政策
        for i, (policy, original_reason) in enumerate(failed_policies, 1):
            if self.stop_requested:
                if callback:
                    callback("停止重试")
                logger.info("[停止] 停止重试")
                break

            if callback:
                callback(f"\n重试进度: [{i}/{len(failed_policies)}]")
            logger.info(f"重试进度: [{i}/{len(failed_policies)}]")

            self.progress.current_policy_id = policy.id
            self.progress.current_policy_title = policy.title
            self._update_progress()

            if callback:
                callback(f"重试政策: {policy.title}")
                callback(f"原失败原因: {original_reason}")

            result_policy = self.crawl_single_policy(policy, callback)

            if result_policy:
                self.progress.completed_count += 1
                self.progress.completed_policies.append(policy.id)
                if callback:
                    callback("   ✓ 重试成功")
            else:
                self.progress.failed_count += 1
                failed_policy_info = {
                    "id": policy.id,
                    "title": policy.title,
                    "link": policy.link or policy.source or policy.url,
                    "pub_date": policy.pub_date,
                    "doc_number": policy.doc_number,
                    "reason": f"重试失败: {original_reason}",
                }
                self.progress.failed_policies.append(failed_policy_info)
                if callback:
                    callback("   ✗ 重试仍然失败")

            self._update_progress()
            time.sleep(self.config.get("request_delay", 2))

        self.progress.end_time = datetime.now(timezone.utc)
        self._update_progress()

        if callback:
            callback("\n重试完成")
        logger.info("重试完成")
        logger.info(f"总计: {self.progress.total_count} 条")
        logger.info(f"成功: {self.progress.completed_count} 条")
        logger.info(f"失败: {self.progress.failed_count} 条")
        logger.info(f"成功率: {self.progress.success_rate:.2f}%")

        return self.progress

    def close(self):
        """关闭爬虫"""
        if hasattr(self.api_client, "close"):
            self.api_client.close()
